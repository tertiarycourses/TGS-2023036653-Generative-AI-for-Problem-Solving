#!/usr/bin/env python3
"""Generate the WSQ Generative AI for Problem Solving Lesson Plan (LP) DOCX.

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 2-day schedule tables (9:30am-6:30pm, 8 training hours/day, 1h
lunch, tea within, final assessment Day 2). Topics/activities come from
course_data + the domain data files so the LP stays aligned with the deck,
guide and activities.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
ACT = DOMAIN1 + DOMAIN2 + DOMAIN3
import prodoc


def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "activities")):
            return d
    return os.path.dirname(os.path.dirname(HERE))


REPO = _find_repo(HERE); ASSETS = os.path.join(os.path.dirname(HERE), "assets")

BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
HEADER_FILL = "1F6FEB"; TOPIC_FILL = "E8F0FE"; BREAK_FILL = "FFF4E5"
LUNCH_FILL = "FDE9D9"; ASSESS_FILL = "E8F7EE"

# Slide anchors are DERIVED FROM THE BUILT DECK by derive_slide_map.py, which the
# orchestrator runs after the slides and before this file. Hand-maintained numbers
# drift the moment a slide is added or removed and send the trainer to the wrong
# slide at every transition, so they are never hardcoded here.
try:
    from slide_map import ANCHORS
except ImportError:
    raise SystemExit("slide_map.py missing — run derive_slide_map.py after building the deck.")


def rng(pair):
    lo, hi = pair
    return "Slide %d" % lo if lo == hi else "Slides %d-%d" % (lo, hi)


def sl(n):
    """Slide reference for activity n, straight from the deck."""
    return rng(ANCHORS["activities"][n])


def one(key, idx=0):
    v = ANCHORS[key]
    return "Slide %d" % (v[idx] if isinstance(v, list) else v)


def span(a, b):
    return "Slides %d-%d" % (a, b)


def act(n):
    a = [x for x in ACT if x["num"] == n][0]
    return "Activity %d: %s (%s)" % (a["num"], a["title"].split("— ", 1)[-1], sl(n))


A = ANCHORS
SLIDES = {n: sl(n) for n in A["activities"]}


# (start, end, minutes, kind, activity_text)
SCHEDULE = {
    1: (C.DAY_THEMES[1], [
        ("9:30", "10:00", 30, "admin",
         "Welcome, trainer and learner introductions, ground rules, learning outcomes, course outline "
         "and mandatory digital attendance (AM). " + span(1, A["topics"][1] - 1)),
        ("10:00", "11:00", 60, "topic",
         "Topic 1 — Problem framing: symptom vs problem vs root cause, well-defined vs ill-defined problems, "
         "the IDEAL cycle, the six elements of a workplace problem statement, cognitive barriers, and how "
         "GenAI fits the problem-solving cycle. " + span(A["topics"][1], A["activities"][1][0] - 1)),
        ("11:00", "11:15", 15, "break", "Tea break (within training time)"),
        ("11:15", "12:30", 75, "lab", "Hands-on: " + act(1)),
        ("12:30", "13:30", 60, "lunch", "Lunch break"),
        ("13:30", "14:00", 30, "topic",
         "Digital attendance (PM). Topic 1 continued — the root cause toolkit: choosing between 5 Whys, "
         "Fishbone, Pareto and System Loops. "
         + span(A["activities"][1][1] + 1, A["activities"][2][0] - 1)),
        ("14:00", "15:00", 60, "lab", "Hands-on: " + act(2)),
        ("15:00", "16:00", 60, "lab", "Hands-on: " + act(3)),
        ("16:00", "16:15", 15, "break", "Tea break (within training time)"),
        ("16:15", "17:15", 60, "lab", "Hands-on: " + act(4)),
        ("17:15", "18:15", 60, "lab", "Hands-on: " + act(5)),
        ("18:15", "18:30", 15, "recap", "Day 1 recap, Q&A and PM digital attendance. " + one("recap_day1")),
    ]),
    2: (C.DAY_THEMES[2], [
        ("9:30", "10:00", 30, "recap",
         "Day 1 recap and mandatory digital attendance (AM). " + span(A["end_day1"], A["topics"][2] - 1)),
        ("10:00", "10:45", 45, "topic",
         "Topic 2 — Divergence before convergence: the solution funnel, divergent techniques (brainstorming, "
         "SCAMPER, Generic Parts, analogy, nominal group), and separating hard constraints from habits. "
         + span(A["topics"][2], A["activities"][6][0] - 1)),
        ("10:45", "11:00", 15, "break", "Tea break (within training time)"),
        ("11:00", "12:00", 60, "lab", "Hands-on: " + act(6)),
        ("12:00", "12:30", 30, "topic",
         "Convergence tools: the Impact-Ease matrix, the weighted decision matrix and sensitivity testing. "
         + span(A["activities"][6][1] + 1, A["activities"][7][0] - 1)),
        ("12:30", "13:30", 60, "lunch", "Lunch break"),
        ("13:30", "14:20", 50, "lab", "Digital attendance (PM). Hands-on: " + act(7)),
        ("14:20", "15:10", 50, "lab", "Hands-on: " + act(8)),
        ("15:10", "15:40", 30, "topic",
         "Topic 3 — Implementation and change: factors affecting effectiveness, stakeholder resistance, "
         "change management actions and the messenger principle. " + span(A["topics"][3], A["activities"][9][0] - 1)),
        ("15:40", "15:55", 15, "break", "Tea break (within training time)"),
        ("15:55", "16:45", 50, "lab", "Hands-on: " + act(9)),
        ("16:45", "17:30", 45, "lab", "Hands-on: " + act(10) + ". Sustaining the gain (PDCA). " + span(A["activities"][10][1] + 1, A["summary"])),
        ("17:30", "17:45", 15, "recap",
         "Course summary, Q&A, course feedback and mandatory TRAQOM survey. " + span(A["summary"] + 1, A["briefing"][-1] - 1)),
        ("17:45", "18:00", 15, "assess", "Briefing for Assessment. " + one("briefing", -1)),
        ("18:00", "18:30", 30, "assess",
         "Digital attendance (Assessment). Final Assessment — Written Assessment (WA, short-answer questions) "
         "and Case Study (CS), open book. " + span(A["assessment"][-1], A["attendance"][-1])),
    ]),
}

# ------------------------------------------------ build document
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc, "LESSON PLAN", C.TITLE, C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc, [
    ("20.0", "1 March 2025",
     "Legacy master lesson plan — Innovative Problem Solving with Generative AI.", C.TRAINER),
    (C.VERSION.lstrip("v"), C.VERSION_DATE,
     "Major revision. Retitled to the published course title. Rebuilt around 10 real-life workplace case-study "
     "activities (ShopFront SG, Meridian Health, Horizon Bank, Nexa Logistics) each with scenario, evidence, "
     "discussion questions and debrief. Added the four problem-solving ed-tools (5 Whys, Fishbone, Pareto, "
     "System Loop), a 10-prompt GenAI library, and content beefed up from Skills for Success, IMD, MIT CCMIT, "
     "Six Sigma, LSE Business Review and current GenAI problem-solving research.", C.TRAINER),
])
prodoc.add_toc(doc)


def H(text, level=1):
    return doc.add_heading(text, level=level)


H("Course Information", 1)
info = [("Course Title", C.TITLE),
        ("WSQ Course Reference", C.COURSE_CODE),
        ("TSC Alignment", "%s (%s)" % (C.TSC_TITLE, C.TSC_CODE)),
        ("Training Provider", C.ORG + "  (" + C.UEN.replace('UEN: ', 'UEN ') + ")"),
        ("Duration", "2 days · 8 training hours per day (16 hours)"),
        ("Daily Timing", "9:30 am – 6:30 pm (1-hour lunch; tea breaks within training time)"),
        ("Mode", "Instructor-led, case-study based. Physical classroom, synchronous ZOOM or corporate on-site"),
        ("Class Size", "Maximum 20 learners; activities run in teams of 3-5"),
        ("Trainer", C.TRAINER)]
t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
for k, v in info:
    c = t.add_row().cells
    c[0].text = ""; r = c[0].paragraphs[0].add_run(k); r.bold = True; r.font.size = Pt(10)
    prodoc._shade_cell(c[0], TOPIC_FILL)
    c[1].text = ""; c[1].paragraphs[0].add_run(v).font.size = Pt(10)

H("Learning Outcomes", 1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size = Pt(10.5)

H("Target Audience and Prerequisites", 1)
for a in ["Operations, project, HR, quality, process-improvement and team-leadership roles that must "
          "diagnose and resolve recurring workplace problems.",
          "Beginner level — no prior problem-solving or AI training is required.",
          "Minimum 3 GCE 'O' Level passes including English, or WPL Level 5.",
          "At least 1 year of working experience.",
          "A laptop with a browser and access to a generative AI tool (ChatGPT, Copilot, Gemini or Claude)."]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size = Pt(10.5)

H("Training Resources", 1)
for a in ["Trainer slide deck (%s), Learner Guide and activity briefs on %s" % (C.VERSION, C.LMS),
          "Problem-solving ed-tools: " + "; ".join("%s (%s)" % (e["name"], e["url"]) for e in C.EDTOOLS),
          "Generative AI tools: " + C.GENAI_TOOLS,
          "Flip charts, post-it notes and markers for team activities"]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size = Pt(10.5)

H("Assessment", 1)
for a in [C.ASSESSMENT["written"], C.ASSESSMENT["practical"],
          "Format: " + C.ASSESSMENT["openbook"],
          "The final assessment is conducted on Day 2 from 5:45 pm, after the Briefing for Assessment.",
          C.ASSESSMENT["note"]]:
    p = doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size = Pt(10.5)


def set_cell(cell, text, bold=False, size=9.5, color=None, fill=None, align=None):
    cell.text = ""; p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = "Arial"
    if color: r.font.color.rgb = color
    if fill: prodoc._shade_cell(cell, fill)


KIND_FILL = {"topic": TOPIC_FILL, "break": BREAK_FILL, "lunch": LUNCH_FILL,
             "assess": ASSESS_FILL, "admin": "F3F5F8", "recap": "F3F5F8", "lab": None}

H("Course Schedule", 1)
for day, (theme, rows) in SCHEDULE.items():
    H("Day %d — %s" % (day, theme), 2)
    tbl = doc.add_table(rows=0, cols=3); tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.add_row().cells
    for i, htext in enumerate(["Time", "Duration", "Topic / Activity"]):
        set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
    training = 0
    for start, end, mins, kind, text in rows:
        cells = tbl.add_row().cells; fill = KIND_FILL.get(kind)
        set_cell(cells[0], "%s–%s" % (start, end), bold=(kind in ("topic", "assess")), size=9.5, fill=fill)
        set_cell(cells[1], "%d min" % mins, size=9.5, fill=fill)
        set_cell(cells[2], text, bold=(kind in ("topic", "assess")), size=9.5, fill=fill)
        if kind != "lunch": training += mins
    for row in tbl.rows:
        row.cells[0].width = Inches(1.15); row.cells[1].width = Inches(0.9); row.cells[2].width = Inches(4.75)
    p = doc.add_paragraph()
    r = p.add_run("Total training time: %d minutes (%d hours)." % (training, training // 60))
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
    assert training == 480, "Day %d training minutes = %d, expected 480" % (day, training)

H("Activity Reference (aligned to TSC abilities)", 1)
tt = doc.add_table(rows=0, cols=4); tt.style = "Table Grid"
hdr = tt.add_row().cells
for i, htext in enumerate(["Topic", "TSC coverage", "Activities", "Case study"]):
    set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for tp in C.TOPICS:
    acts = [a for a in ACT if a["topic"] == tp["num"]]
    cells = tt.add_row().cells
    set_cell(cells[0], "Topic %s: %s" % (tp["code"], tp["title"]), bold=True, size=9, fill=TOPIC_FILL)
    set_cell(cells[1], tp["weighting"], size=9, fill=TOPIC_FILL)
    set_cell(cells[2], ", ".join("Activity %d" % a["num"] for a in acts), size=9)
    set_cell(cells[3], "; ".join(sorted({a["case_title"].split(" — ")[0] for a in acts})), size=9)

H("Detailed Activity Schedule", 1)
at = doc.add_table(rows=0, cols=5); at.style = "Table Grid"
hdr = at.add_row().cells
for i, htext in enumerate(["#", "Activity", "Duration", "Grouping", "Slides"]):
    set_cell(hdr[i], htext, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF), fill=HEADER_FILL)
for a in ACT:
    cells = at.add_row().cells
    set_cell(cells[0], str(a["num"]), bold=True, size=9, fill=TOPIC_FILL)
    set_cell(cells[1], a["title"].split("— ", 1)[-1], size=9)
    set_cell(cells[2], a["duration"], size=9)
    set_cell(cells[3], a["grouping"], size=9)
    set_cell(cells[4], SLIDES.get(a["num"], ""), size=9)
for row in at.rows:
    row.cells[0].width = Inches(0.35); row.cells[1].width = Inches(3.1)
    row.cells[2].width = Inches(0.9); row.cells[3].width = Inches(0.9); row.cells[4].width = Inches(1.05)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT = os.path.join(REPO, "courseware", "LP-%s.docx" % C.SHORT_TITLE)
doc.save(OUT)
print("Saved", OUT)
