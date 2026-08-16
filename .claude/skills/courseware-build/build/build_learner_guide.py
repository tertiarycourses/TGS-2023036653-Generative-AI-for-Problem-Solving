#!/usr/bin/env python3
"""Generate the WSQ Generative AI for Problem Solving Learner Guide as BOTH a
Markdown mirror (LG-*.md at repo root) and a DOCX (courseware/LG-*.docx) from one
source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per activity (Objective · Case scenario · Evidence · Step-by-step ·
Prompt · Discussion questions · Debrief · Self-check), plus the theory chapters,
the prompt library and a glossary. All content is driven by course_data + the
domain data files, keeping the LG 100% aligned with the deck, LP and activities.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
ACT = DOMAIN1 + DOMAIN2 + DOMAIN3
import prodoc


def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "activities")):
            return d
    return os.path.dirname(os.path.dirname(HERE))


REPO = _find_repo(HERE); ASSETS = os.path.join(os.path.dirname(HERE), "assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B = []
def h1(t): B.append(("h1", t))
def h2(t): B.append(("h2", t))
def h3(t): B.append(("h3", t))
def p(t):  B.append(("p", t))
def bullets(xs): B.append(("bullets", xs))
def steps(xs): B.append(("steps", xs))
def code(t): B.append(("code", t))
def note(t): B.append(("note", t))
def table(hdr, rows): B.append(("table", hdr, rows))
def numbered(xs): B.append(("numbered", xs))
def rule(): B.append(("rule",))


# ================================================================ INTRODUCTION
h1("Introduction")
p("This Learner Guide accompanies the WSQ course %s (%s), conducted by %s. It is your "
  "step-by-step companion for all ten hands-on case-study activities, and your reference after the "
  "course when you take these methods back to your own workplace." % (C.TITLE, C.COURSE_CODE, C.ORG))
p("The course is built around a single idea: most organisations are very good at generating "
  "solutions and very poor at defining problems. The result is expensive activity that changes "
  "nothing. Over two days you will learn a complete, repeatable method — frame the problem, diagnose "
  "the true root cause, generate a wide solution set, converge on a defensible choice, implement it "
  "against real human resistance, and prove whether it worked.")
p("Generative AI is used throughout, but never as an oracle. It is a divergence engine and a "
  "tireless challenger: it expands what you consider and accelerates how fast you get there. The "
  "evidence, the judgement and the accountability remain yours. Every activity therefore asks you to "
  "compare what the AI produced against what your evidence actually supports.")
p("This course aligns to the SkillsFuture Singapore Skills Framework Technical Skill and Competency "
  "\"%s\" (%s)." % (C.TSC_TITLE, C.TSC_CODE))

h2("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h2("TSC Abilities Covered")
bullets(C.TSC_ABILITIES)

h2("TSC Knowledge Covered")
bullets(C.TSC_KNOWLEDGE)

h2("How to Use This Guide")
bullets([
    "Each activity section is self-contained: the scenario, the evidence, the exact steps, the GenAI "
    "prompt, the discussion questions and the debrief.",
    "Work through the steps in order during the class. The steps are deliberately detailed so you can "
    "repeat the method at work without the trainer present.",
    "The prompts are written to be copied verbatim into ChatGPT, Microsoft Copilot, Google Gemini or "
    "Claude. Replace anything in <<double angle brackets>> with your own situation.",
    "The Self-check at the end of each activity tells you when your output is genuinely finished, "
    "rather than merely complete-looking.",
    "The assessment is open book: this guide and the slides are permitted. Internet search and AI "
    "tools are NOT permitted during assessment.",
])

h2("Tools You Will Use")
table(["Tool", "URL", "What it is for"],
      [[e["name"], e["url"], e["use"]] for e in C.EDTOOLS] +
      [["Generative AI", "ChatGPT / Copilot / Gemini / Claude",
        "Framing, cause hypotheses, ideation, scoring and structuring evaluation."]])

# ================================================================ TOPIC 1 THEORY
h1("Topic 1 — Identifying Performance Gaps and Root Causes")

h2("1.1 Why Problem Definition Comes First")
p("Charles Kettering's line — \"a problem well stated is a problem half solved\" — is the most "
  "under-applied principle in management. A vague problem statement does not merely slow you down; "
  "it actively misdirects spending, because every department resolves the ambiguity in favour of the "
  "work it already wanted to do.")
p("In the ShopFront SG case you will meet in Activity 1, one sentence — \"customers are unhappy and "
  "sales are dropping\" — sends Marketing to a discount campaign, IT to an app redesign and Customer "
  "Service to hiring temps. Ninety days and S$180,000 later, nothing has changed. Not one of those "
  "three teams did anything unreasonable. The failure was upstream of all of them.")

h3("Symptom, problem and root cause")
table(["", "Symptom", "Problem", "Root cause"],
      [["What it is", "What people notice and complain about",
        "The measurable performance deficiency", "The system condition that produces it"],
       ["ShopFront example", "\"Customers are unhappy\"",
        "Checkout completion fell from 68% to 51%", "No capacity planning informed by analytics"],
       ["Effect of fixing it", "Feels responsive; changes nothing",
        "Moves the number this quarter", "Stops the problem recurring"],
       ["Typical response", "A campaign or a memo", "A project", "A change to policy, process or incentive"]])
p("Treating a symptom always feels faster, which is precisely why it is the most expensive habit in "
  "problem solving. The discipline this course teaches is to spend deliberate, uncomfortable time at "
  "the problem and root-cause layers before authorising any spend.")

h3("Well-defined and ill-defined problems")
p("A well-defined problem has a clear goal and a known method — \"reduce app load time to under two "
  "seconds\". An ill-defined problem has neither — \"improve the customer experience\". Most problems "
  "arrive at your desk ill-defined, which means framing is not preparation for the work; framing IS "
  "the work. Generative AI is particularly valuable here because it can rapidly restate an ill-defined "
  "problem in several different structured forms, letting you choose the framing that best fits your "
  "evidence.")

h3("The IDEAL cycle")
p("The IDEAL heuristic (Bransford & Stein, taught at MIT) gives you a five-stage backbone that this "
  "entire course follows:")
numbered([
    "Identify the problem — recognise that a problem exists and pinpoint what it actually is, by "
    "asking questions and gathering information rather than assuming.",
    "Define the context — investigate the facts. When did this last work correctly, and what changed?",
    "Explore possible strategies — brainstorm and develop multiple solution pathways before committing.",
    "Act on the best solution — implement the selected approach.",
    "Look back and learn — evaluate the outcome and document the lesson.",
])
note("Most teams start at step 4. The cost of skipping steps 1 and 2 is always paid later, with "
     "interest, because you fix the wrong thing and then have to fix it again.")

h3("The six sub-skills of problem solving")
p("The Government of Canada's Skills for Success framework defines problem solving as the ability to "
  "identify, analyse, propose solutions and make decisions. It breaks into six sub-skills, which map "
  "directly onto this course:")
table(["Sub-skill", "What it means", "Where you practise it"],
      [["Issue identification", "Is the problem familiar or novel, simple or complex?", "Activity 1"],
       ["Information gathering", "Collect facts; find how similar problems were solved before", "Activities 1-4"],
       ["Analysis", "Break the problem into parts; establish cause and effect", "Activities 2-5"],
       ["Action generation", "Develop several solutions across short and long horizons", "Activities 6-7"],
       ["Implementation", "Select and execute, staying flexible", "Activities 8-9"],
       ["Evaluation", "Assess effectiveness and extract the lesson", "Activity 10"]])

h2("1.2 Writing a Measurable Problem Statement")
p("A workplace problem statement is not a description of a feeling. It is a contract that lets a "
  "team start work, lets a sponsor approve spend, and lets everyone evaluate the outcome later. It "
  "carries six elements, and all six are compulsory.")
table(["Element", "What it captures", "Failure if missing"],
      [["Context", "Who, where, and which process is affected", "Scope creep; wrong people involved"],
       ["Metric", "The single primary measure you will move", "Teams optimise different things"],
       ["Baseline", "Where the metric is today, with the number", "You can never prove improvement"],
       ["Target", "Where it must reach, with the number", "No definition of success"],
       ["Timeframe", "By when", "Nothing is ever late"],
       ["Constraints", "Budget, headcount, compliance, brand", "Solutions proposed that cannot be adopted"]])
p("The Baseline element deserves particular attention. It is the element teams most often skip, and "
  "it is the one that makes evaluation possible sixteen hours later in Topic 3. If you take one "
  "habit from this course, take this: never accept a problem statement without a baseline number.")

h3("Worked example")
p("Weak: \"Our customers are not happy with our services.\"")
p("Strong: \"Checkout completion on the ShopFront SG app and web store has fallen from 68% to 51% "
  "(benchmark 70%) over 12 months, alongside app load times rising from 2.1s to 4.8s, contributing "
  "to a S$0.9m per month decline in online sales. We aim to restore checkout completion to at least "
  "68% and app load to under 2.0s within 90 days, without additional headcount and without changes "
  "requiring PDPA re-consent.\"")

h2("1.3 Cognitive Barriers")
p("Problem solving fails for predictable psychological reasons as often as for analytical ones. Know "
  "these six and you will catch yourself in the act:")
table(["Barrier", "How it shows up at work"],
      [["Functional fixedness", "Seeing a resource only for its usual purpose, so obvious re-uses stay invisible"],
       ["Mental set", "Re-using the solution that worked last time on a problem that has since changed"],
       ["Confirmation bias", "Collecting only the evidence that supports the cause you already suspect"],
       ["Unnecessary constraints", "Treating habits as rules; most 'fixed' constraints are never enforced by anyone"],
       ["Irrelevant information", "Drowning in data while the three numbers that matter go unexamined"],
       ["Consensus fatigue", "Accepting a weak option because the meeting has gone on too long"]])
note("Two of these — cognitive fixedness and consensus fatigue — are specifically named by IMD as the "
     "biases that most reduce the quality of leadership problem solving. Activity 6 is designed to "
     "break the first; Activity 7 is designed to defeat the second.")

h2("1.4 Generative AI in Problem Solving")
p("Generative AI does not solve your problem. It expands what you consider and accelerates how fast "
  "you get there. Understanding exactly where it is strong and where it is weak is the difference "
  "between a productivity gain and a confidently-wrong decision.")
table(["Dimension", "GenAI is strong", "GenAI is weak"],
      [["Breadth of options", "Generates 20 ideas in seconds, free of office politics", "Cannot tell which are real"],
       ["Framing", "Restates vague problems into structured, testable form", "Invents specifics your evidence never supplied"],
       ["Analogy", "Borrows working patterns across every industry it has read", "May transfer a pattern that does not apply"],
       ["Arithmetic", "Improving, and good at showing method", "Still makes errors — verify every number"],
       ["Context", "Applies general best practice competently", "Does not know your people, history or politics"],
       ["Accountability", "Tireless, patient, unbiased by hierarchy", "Cannot be held responsible for the decision"]])

h3("Why GenAI hallucinates specifics — and what to do about it")
p("A language model is trained to produce plausible continuations. When your prompt lacks a detail "
  "the answer seems to require, the model supplies a plausible one rather than leaving a gap. This is "
  "not a malfunction; it is the mechanism working as designed. In problem solving it is dangerous "
  "because invented specifics look identical to established facts.")
p("The defence is procedural, not technical. Always supply your real evidence in the prompt, and "
  "always instruct the model to mark what it assumed. You will practise this in every single "
  "activity, and it is the most valuable AI habit you will take away from this course.")

h3("Known limitations from current research")
bullets([
    "Poor mathematical reliability without external tools — always check the arithmetic yourself.",
    "Weak memory, planning and reasoning in single-agent use — break complex problems into steps.",
    "Unreliable on domain-specific questions without access to your proprietary data.",
    "Retrieval-augmented generation (RAG) and tool use materially improve reliability, which is why "
    "enterprise deployments connect models to real company data rather than relying on training data.",
])

h3("Prompting habits that work")
bullets([
    "Give it a role — \"Act as a root cause analysis expert\" sets the reasoning frame.",
    "Supply the evidence — paste your real numbers; without them the model invents plausible ones.",
    "Demand structure — name the fields you want back, so outputs are comparable across teams.",
    "Ask it to flag assumptions — \"mark anything you assumed that my evidence did not support\".",
    "Make it challenge you — \"ask me for evidence rather than accepting my answer\".",
    "Iterate — the first output is a draft. The third is usually useful.",
])

h2("1.5 The Root Cause Toolkit")
p("Four tools, four different jobs. Mature problem solvers use them together rather than choosing "
  "between them:")
table(["Tool", "What it gives you", "Use it when", "Limitation"],
      [["5 Whys", "Depth on one causal chain", "The problem is focused and singular", "Follows only one branch"],
       ["Fishbone", "Breadth across all dimensions", "Causes are many and unclear", "Shallow on each branch"],
       ["Pareto", "Priority by contribution", "You have counted, categorised data", "Ranks frequency, not severity"],
       ["System Loops", "Feedback dynamics over time", "Fixes keep rebounding", "Needs time-series insight"]])
note("The tools chain naturally: Pareto tells you WHERE to look, 5 Whys tells you WHY it happens, "
     "Fishbone makes sure you have not missed a dimension, and System Loops explains why your last "
     "fix stopped working.")

# ================================================================ TOPIC 2 THEORY
h1("Topic 2 — Developing Corrective Plans and Evaluating Solutions")

h2("2.1 Divergence Before Convergence")
p("Generating and judging are different cognitive modes, and running them at the same time destroys "
  "your best ideas. A team that evaluates while it generates typically stops at six to eight "
  "conventional options, because every unusual idea is killed by the first objection before it can "
  "be developed.")
p("The solution funnel therefore runs in strict sequence: diverge widely, cluster by theme, screen "
  "quickly, score rigorously, then commit. Each stage has its own tool, and mixing them is the most "
  "common facilitation error in workplace problem solving.")

h2("2.2 Techniques for Generating Solutions")
table(["Technique", "How it works", "Best for"],
      [["Brainstorming", "All ideas welcome, judgement suspended, quantity first", "General divergence"],
       ["SCAMPER", "Seven prompts applied to a fixed assumption", "Breaking cognitive fixedness"],
       ["Generic Parts Technique", "Strip the problem to abstract functions", "Finding solutions in other domains"],
       ["Analogy / cross-industry", "Borrow a working pattern from another sector", "Novel problems"],
       ["Reverse brainstorming", "Ask how to make it worse, then invert", "Stuck teams"],
       ["Nominal group technique", "Silent individual generation, then group discussion", "Avoiding groupthink"],
       ["Six Thinking Hats", "Examine from facts, feelings, risks, benefits, creativity, process", "Balanced review"],
       ["Mind mapping", "Radiate ideas from the central problem", "Seeing relationships"]])

h3("SCAMPER in detail")
table(["Move", "The question it asks"],
      [["Substitute", "What can be swapped out — different person, material, place or rule?"],
       ["Combine", "What can be merged — two steps, two roles, two services?"],
       ["Adapt", "What works elsewhere that we could borrow?"],
       ["Modify", "What if we magnified, shrank or reshaped it?"],
       ["Put to another use", "What else could this resource or step do?"],
       ["Eliminate", "What if we simply removed it? (Often the biggest unlock.)"],
       ["Reverse", "What if we ran it backwards, or swapped who does what?"]])

h3("Hard constraints versus habits")
p("The single most valuable output of a SCAMPER session is discovering that a constraint everyone "
  "treated as fixed is actually a habit. In the Meridian Health case, the MOH phlebotomy competency "
  "standard is a genuine hard constraint — it is a regulation with an enforcing body and a penalty. "
  "The 07:00-09:30 draw window at the centre is a habit; it feels equally fixed, but it exists only "
  "because nobody has questioned it.")
p("The test is simple and you should apply it every time: who enforces this, and what is the penalty "
  "for breaking it? If you cannot name an enforcer, you are looking at a habit.")

h2("2.3 Converging on a Decision")
p("Convergence needs explicit criteria, or the decision defaults to whoever is most senior or most "
  "persistent. Use two tools in sequence — a fast screen, then a rigorous instrument.")

h3("The Impact-Ease matrix")
table(["Quadrant", "Meaning", "What to do"],
      [["Quick Win", "High impact, easy to implement", "Do these first — they build momentum and credibility"],
       ["Big Bet", "High impact, hard to implement", "Do these, but plan and resource them properly"],
       ["Fill-In", "Low impact, easy to implement", "Do only if you have spare capacity"],
       ["Thankless Task", "Low impact, hard to implement", "Do not start these"]])

h3("The weighted decision matrix")
p("The Impact-Ease screen is directional. When you must defend a decision to a board, you need a "
  "weighted matrix: explicit criteria, explicit weights, scores from 1 to 5, and a computed total.")
p("The critical facilitation move is to agree the WEIGHTS before scoring anything. Most arguments "
  "that appear to be about scores are really disagreements about weights, and surfacing that early "
  "converts a political argument into an explicit, recorded choice.")
note("Always run a sensitivity test. Change one weight materially and re-rank. If the winner changes, "
     "your recommendation is fragile and you must say so. If it holds across several weightings, you "
     "have a robust recommendation — and boards trust robust recommendations.")

h2("2.4 The Corrective Action Plan")
p("A corrective action plan converts a chosen solution into something a colleague could execute "
  "without you in the room. Eight components, each of which exists because its absence causes a "
  "specific, predictable failure:")
table(["Component", "Why it exists", "Failure if missing"],
      [["Root cause addressed", "Ties the action to the diagnosis", "Solution drifts back to treating a symptom"],
       ["Corrective action", "The specific thing being done", "Vague intent that never becomes execution"],
       ["Owner (named role)", "Single point of accountability", "Everyone assumes someone else is doing it"],
       ["Timeline / milestone", "Makes progress checkable", "Slips indefinitely and nobody notices"],
       ["Resources required", "Budget, people and tools committed", "Stalls at the first resource conflict"],
       ["Success measure", "Defines what 'worked' means", "Cannot evaluate; endless debate"],
       ["Risk and mitigation", "Anticipates what breaks it", "Surprised by an entirely predictable failure"],
       ["Review checkpoint", "Forces a continue/stop decision", "Runs on long past the point of failure"]])
note("An owner must be a named role, never a department. 'Operations' cannot be phoned, cannot be "
     "held to a date and cannot be asked why something slipped. 'Centre Operations Manager' can.")

# ================================================================ TOPIC 3 THEORY
h1("Topic 3 — Selecting, Implementing and Measuring Effectiveness")

h2("3.1 Why Good Solutions Fail")
p("Most solutions fail on adoption rather than on design. The organisation in Activity 9 bought a "
  "S$400,000 CRM that worked perfectly and that nobody used. The technology was not the problem. "
  "This is why the Skills Framework treats 'factors affecting the effectiveness of an implementation "
  "plan' (K8) as a distinct body of knowledge: implementation is mostly about people.")
table(["Factor", "What good looks like"],
      [["Sponsorship", "A senior owner who can re-prioritise competing work, not merely endorse yours"],
       ["Sequencing", "Dependencies respected — never drive demand into a system you have not yet fixed"],
       ["Capacity", "Teams already at capacity are given real relief, not encouragement"],
       ["Stakeholder buy-in", "Resistance treated as a design input, surfaced early"],
       ["Communication", "Regular, specific, and continuing past launch"],
       ["Measurement", "Window, baseline and control defined at day zero"]])

h2("3.2 Reading and Working With Resistance")
p("Resistance is information. It usually tells you something true about your plan that you did not "
  "know. Learning to read what someone is actually objecting to — as opposed to what they say — is "
  "the core skill of implementation.")
table(["Stakeholder type", "What they actually object to", "The move that works"],
      [["At-capacity team", "The workload, not the idea", "Real relief — re-prioritise other work or fund capacity"],
       ["Face at risk", "Being shown to have been wrong", "Reframe rather than relitigate; give them ownership of a win"],
       ["External party", "Cost and scrutiny", "Co-design the standard, then tie it to commercial consequence"],
       ["Job-security fear", "A rational threat to their livelihood", "Commit only to what you can actually guarantee"],
       ["The indifferent", "Nothing at all — which is the danger", "Give them a stake before you need them"]])
note("The messenger matters as much as the message. The same words land completely differently coming "
     "from the sponsor, a peer, or the analyst who produced the diagnosis. Assign a messenger to every "
     "difficult conversation and be able to justify the choice.")
p("Indifference kills more rollouts than opposition does, because nobody escalates an absence of "
  "enthusiasm. Opposition is visible and gets managed; indifference is invisible until launch day.")

h3("Change management actions")
bullets([
    "Leadership alignment — communicate the business goal and tie measurable KPIs to outcomes.",
    "Quick wins first — deliver an early visible improvement to build momentum and credibility.",
    "Involve affected teams early, in solution design rather than after the decision.",
    "Co-design with likely resisters — an imposed standard is complied with on paper; a co-designed "
    "one is actually adopted.",
    "Incentives and accountability — link the change to metrics people are genuinely measured on.",
    "Train and support — SOPs, quick guides and a real support channel; adoption dies without them.",
])

h2("3.3 Measuring Effectiveness")
p("\"It's working\" is not an evaluation. A credible evaluation gives a verdict per metric against "
  "both baseline and target, states attribution explicitly, and is honest about what cannot be "
  "determined from the available evidence.")

h3("Leading and lagging indicators")
table(["", "Leading indicators", "Lagging indicators"],
      [["Respond in", "Days to weeks", "Months to quarters"],
       ["Examples", "Page load time, checkout completion, adoption rate", "CSAT, repeat purchase, brand trust"],
       ["They tell you", "Whether the change is taking hold", "Whether the outcome actually improved"],
       ["Risk if used alone", "An early signal that may not persist", "Feedback arrives too late to steer"],
       ["How to use them", "Steer with these week to week", "Set realistic horizons and judge outcomes"]])
p("Expecting a lagging indicator to move inside 90 days is a planning error, not an execution "
  "failure. Customer satisfaction reflects accumulated experience across several purchase cycles — "
  "for a fashion retailer, six to nine months rather than three.")

h3("Attribution and confounding")
p("In a real workplace you rarely get a clean experiment. Something else always happens during your "
  "measurement window — a seasonal peak, a competitor's move, another internal project. Good practice "
  "is to define the measurement window at day zero, hold a control where possible (for example a "
  "staged rollout by region), and segment the data afterwards.")
p("Where you cannot establish attribution, say so plainly. Stating what you cannot determine is what "
  "makes the rest of your evaluation credible to a finance director.")

h3("Solution failure versus implementation failure")
table(["", "Solution failure", "Implementation failure"],
      [["What happened", "It was genuinely done, and it did not work", "It was never really done"],
       ["Evidence to check", "Open rates, click rates, sector benchmarks", "Adoption and compliance rates"],
       ["Right response", "Change or drop the solution", "Fix the rollout and the change management"],
       ["Common error", "Abandoning something that was never adopted", "Blaming people for a bad design"]])
note("This distinction decides where the next tranche of budget goes. Abandoning a solution that was "
     "never actually adopted throws away a fix that might have worked perfectly well.")

h3("Sustaining the gain")
bullets([
    "PDCA / DMAIC-Control — without a control phase, improvements decay back to the old behaviour.",
    "Standardise — write the new way into the SOP, the induction and the checklist.",
    "Monitor — keep the leading indicator on a dashboard someone actually looks at.",
    "Assign ownership — a named role owns the sustained metric after the project team disbands.",
    "Review cadence — a standing checkpoint that carries a decision, not a status update.",
    "Capture the lesson — the 'Look back and learn' step of IDEAL, and the one most often skipped.",
])

# ================================================================ ACTIVITIES
h1("Hands-On Activities")
p("Each activity below is a real-life workplace case study. Work in the team size stated, follow the "
  "steps in order, run the prompt, discuss the questions, and use the debrief to check your thinking. "
  "The Self-check tells you when your output is genuinely finished.")

for a in ACT:
    rule()
    h2(a["title"])
    h3("Objective")
    p(a["objective"])
    h3("At a glance")
    table(["Field", "Detail"],
          [["Case study", a["case_title"]],
           ["Duration", a["duration"]],
           ["Grouping", a["grouping"]],
           ["Tools", a["services"]],
           ["Ed-tool", "%s — %s" % (a["edtool"]["name"], a["edtool"]["url"])]])

    h3("The scenario")
    for para in [x.strip() for x in a["scenario"].split("\n\n") if x.strip()]:
        p(para)

    h3("The evidence")
    p(a["data"]["caption"])
    table(a["data"]["rows"][0], a["data"]["rows"][1:])

    h3("Step-by-step")
    numbered([s[0] for s in a["steps"]])

    h3("The GenAI prompt")
    p("Copy this prompt into your generative AI tool. Replace anything in <<double angle brackets>> "
      "with your own details.")
    code(a["prompt"])

    h3("Discussion questions")
    numbered(a["questions"])

    h3("Debrief")
    for para in [x.strip() for x in a["debrief"].split("\n\n") if x.strip()]:
        p(para)

    h3("Self-check — is your output finished?")
    p(a["test"])

# ================================================================ PROMPT LIBRARY
rule()
h1("The Prompt Library")
p("These are the ten prompts used across the course, collected for reuse at work. Each is written to "
  "be pasted directly into ChatGPT, Microsoft Copilot, Google Gemini or Claude. Replace anything in "
  "<<double angle brackets>> with your own situation, and always paste your real evidence — without "
  "it the model will invent plausible substitutes.")
for i, a in enumerate(ACT, 1):
    h2("Prompt %d — %s" % (i, a["title"].split("— ", 1)[-1]))
    p("Use for: " + a["objective"])
    code(a["prompt"])

# ================================================================ ASSESSMENT
rule()
h1("Preparing for the Assessment")
h2("Format")
bullets([C.ASSESSMENT["written"], C.ASSESSMENT["practical"],
         C.ASSESSMENT["openbook"], C.ASSESSMENT["note"]])

h2("What you should be able to do")
bullets([
    "Write a six-element problem statement from a vague complaint plus evidence, and identify which "
    "metric represents the real performance deficiency.",
    "Run a 5 Whys chain to a systemic root cause and justify each level with evidence.",
    "Populate a fishbone across six dimensions and identify which dimension an intervention targets.",
    "Compute a Pareto cumulative percentage, identify the vital few, and state what Pareto cannot tell you.",
    "Identify reinforcing and balancing loops and explain why a fix rebounded.",
    "Generate a wide solution set and distinguish a hard constraint from a habit.",
    "Apply an Impact-Ease screen and a weighted decision matrix, including a sensitivity test.",
    "Write a corrective action plan with all eight components and a named owner per action.",
    "Predict stakeholder resistance and specify the move and the messenger for each.",
    "Evaluate results against baseline and target, state attribution honestly, and distinguish "
    "solution failure from implementation failure.",
])

h2("Common mistakes to avoid")
bullets([
    "Writing a problem statement with no baseline number — this makes evaluation impossible later.",
    "Stopping a 5 Whys chain at a person ('the relief staff are slow') rather than a system condition.",
    "Leaving a fishbone bone empty and treating that as 'no causes' rather than 'no visibility'.",
    "Claiming a Pareto analysis proves an SLA is achievable without doing the arithmetic.",
    "Proposing 'hire more people' as a leverage point — that is effort, not leverage.",
    "Ranking solutions before the criteria and weights have been agreed and written down.",
    "Naming a department rather than a role as the owner of a corrective action.",
    "Assuming stakeholders will cooperate because the analysis is correct.",
    "Crediting a programme with every improvement while blaming the season for every shortfall.",
])

# ================================================================ GLOSSARY
h1("Glossary")
gl = [
    ("Baseline", "The measured value of a metric before any intervention — the reference point that makes evaluation possible."),
    ("Balancing loop (B)", "A feedback loop that pushes a system toward a target and stabilises it."),
    ("Cognitive fixedness", "The tendency to default to familiar solutions even when the context has changed."),
    ("Confounding", "Another factor changing during your measurement window, making attribution uncertain."),
    ("Consensus fatigue", "Accepting a weak option merely to end a long discussion."),
    ("Corrective action plan (CAP)", "A structured plan converting a chosen solution into owned, timed, measurable actions."),
    ("DMAIC", "Define, Measure, Analyze, Improve, Control — the Six Sigma improvement framework."),
    ("Fishbone (Ishikawa) diagram", "A cause-categorisation diagram organising potential causes across dimensions."),
    ("Five Whys", "Asking 'why' repeatedly, typically five times, to move from symptom to root cause."),
    ("Generic Parts Technique", "Stripping a problem to abstract functional components to break fixedness."),
    ("Hallucination", "A confident, plausible but unsupported output from a generative AI model."),
    ("IDEAL", "Identify, Define, Explore, Act, Look back — a five-stage problem-solving heuristic."),
    ("Ill-defined problem", "A problem with an unclear goal and no standard method — most workplace problems."),
    ("Impact-Ease matrix", "A 2x2 screen classifying options as Quick Wins, Big Bets, Fill-Ins or Thankless Tasks."),
    ("Lagging indicator", "A measure that reflects accumulated outcomes and responds slowly (e.g. CSAT)."),
    ("Leading indicator", "A measure that responds quickly and signals whether a change is taking hold."),
    ("Leverage point", "A place where a small structural change produces a large, self-sustaining effect."),
    ("Pareto analysis", "Ranking causes by contribution to find the vital few responsible for most of the effect."),
    ("PDCA", "Plan-Do-Check-Act — an iterative improvement and control cycle."),
    ("Performance deficiency", "A measurable gap between required and actual performance in a system or process."),
    ("Problem statement", "A structured statement carrying context, metric, baseline, target, timeframe and constraints."),
    ("Prompt engineering", "Structuring instructions to a generative AI model to obtain reliable, usable output."),
    ("RAG", "Retrieval-augmented generation — grounding an AI model in your own data to reduce hallucination."),
    ("Reinforcing loop (R)", "A feedback loop that compounds, making things progressively better or worse."),
    ("Root cause", "The underlying system condition which, if removed, stops the problem recurring."),
    ("SCAMPER", "Substitute, Combine, Adapt, Modify, Put to another use, Eliminate, Reverse."),
    ("Sensitivity test", "Re-ranking options under changed criteria weights to test how robust a decision is."),
    ("Symptom", "The visible effect of a problem, frequently mistaken for the problem itself."),
    ("Systems thinking", "Analysing a problem as a web of interacting feedback loops rather than a linear chain."),
    ("Weighted decision matrix", "A scoring tool applying agreed weights to agreed criteria to rank options defensibly."),
    ("Well-defined problem", "A problem with a clear goal and a known method of solution."),
]
B.append(("dl", gl))


# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)


def render_md():
    out = ["# %s — Learner Guide" % C.TITLE, ""]
    out.append("**WSQ Course Code:** %s  |  **Conducted by:** %s (%s)  |  **Version %s · %s**"
               % (C.COURSE_CODE, C.ORG, C.UEN.replace('UEN: ', 'UEN '), C.VERSION, C.VERSION_DATE))
    out.append("")
    out.append("## Contents"); out.append("")
    for kind, *rest in B:
        if kind == "h1": out.append("- [%s](#%s)" % (rest[0], _anchor(rest[0])))
        elif kind == "h2": out.append("  - [%s](#%s)" % (rest[0], _anchor(rest[0])))
    out.append("")
    for kind, *rest in B:
        if kind == "h1": out += ["", "## %s" % rest[0], ""]
        elif kind == "h2": out += ["", "### %s" % rest[0], ""]
        elif kind == "h3": out += ["**%s**" % rest[0], ""]
        elif kind == "p": out += [rest[0], ""]
        elif kind == "bullets": out += ["- %s" % x for x in rest[0]] + [""]
        elif kind == "numbered":
            out += ["%d. %s" % (i, x) for i, x in enumerate(rest[0], 1)] + [""]
        elif kind == "steps":
            for i, (instr, cmd) in enumerate(rest[0], 1):
                out.append("%d. %s" % (i, instr))
            out.append("")
        elif kind == "table":
            hdr, rows = rest[0], rest[1]
            out.append("| " + " | ".join(hdr) + " |")
            out.append("|" + "|".join(["---"] * len(hdr)) + "|")
            for r in rows:
                out.append("| " + " | ".join(str(c).replace("|", "\\|") for c in r) + " |")
            out.append("")
        elif kind == "code": out += ["```text", rest[0], "```", ""]
        elif kind == "note": out += ["> **Note:** %s" % rest[0], ""]
        elif kind == "rule": out += ["---", ""]
        elif kind == "dl":
            for term, defn in rest[0]: out.append("- **%s** — %s" % (term, defn))
            out.append("")
    return "\n".join(out)


MD_OUT = os.path.join(REPO, "LG-%s.md" % C.SHORT_TITLE)
with open(MD_OUT, "w") as f: f.write(render_md())
print("Saved", MD_OUT)

# ---------------- render DOCX ----------------
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
INKCODE = RGBColor(0x0B, 0x30, 0x60)
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc, "LEARNER GUIDE", C.TITLE, C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc, [
    ("20.0", "1 March 2025",
     "Legacy master learner guide — Innovative Problem Solving with Generative AI.", C.TRAINER),
    (C.VERSION.lstrip("v"), C.VERSION_DATE,
     "Major revision. Retitled to the published course title. Rebuilt around 10 real-life workplace "
     "case-study activities, each with scenario, evidence table, detailed step-by-step, GenAI prompt, "
     "discussion questions, debrief and self-check. Added expanded theory chapters for all three topics, "
     "a reusable 10-prompt library, assessment preparation and a 31-term glossary. Content beefed up "
     "from Skills for Success, IMD, MIT CCMIT, Coursera, Six Sigma, LSE Business Review and current "
     "generative-AI problem-solving research.", C.TRAINER),
])
prodoc.add_toc(doc)


def code_para(text):
    for line in text.split("\n"):
        para = doc.add_paragraph()
        r = para.add_run(line if line else " ")
        r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = INKCODE
        para.paragraph_format.space_after = Pt(0)


def add_table(hdr, rows):
    t = doc.add_table(rows=0, cols=len(hdr)); t.style = "Table Grid"
    cells = t.add_row().cells
    for i, h in enumerate(hdr):
        cells[i].text = ""
        r = cells[i].paragraphs[0].add_run(str(h))
        r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        prodoc._shade_cell(cells[i], "1F6FEB")
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row[:len(hdr)]):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(c))
            r.font.size = Pt(9.5)
            if i == 0: r.bold = True
    doc.add_paragraph()


for kind, *rest in B:
    if kind == "h1": doc.add_heading(rest[0], level=1)
    elif kind == "h2": doc.add_heading(rest[0], level=2)
    elif kind == "h3":
        para = doc.add_paragraph(); r = para.add_run(rest[0])
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BRAND
    elif kind == "p": doc.add_paragraph(rest[0])
    elif kind == "bullets":
        for x in rest[0]: doc.add_paragraph(x, style="List Bullet")
    elif kind == "numbered":
        for x in rest[0]: doc.add_paragraph(x, style="List Number")
    elif kind == "steps":
        for instr, cmd in rest[0]:
            doc.add_paragraph(instr, style="List Number")
    elif kind == "table": add_table(rest[0], rest[1])
    elif kind == "code": code_para(rest[0])
    elif kind == "note":
        para = doc.add_paragraph(); r = para.add_run("Note: ")
        r.bold = True; r.font.color.rgb = BRAND
        para.add_run(rest[0]).font.size = Pt(10)
    elif kind == "rule": doc.add_paragraph("")
    elif kind == "dl":
        for term, defn in rest[0]:
            para = doc.add_paragraph(style="List Bullet")
            r = para.add_run(term + " — "); r.bold = True
            para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT = os.path.join(REPO, "courseware", "LG-%s.docx" % C.SHORT_TITLE)
doc.save(DOCX_OUT)
print("Saved", DOCX_OUT)
