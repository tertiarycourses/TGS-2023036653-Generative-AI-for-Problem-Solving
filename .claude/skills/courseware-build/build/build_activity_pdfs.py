#!/usr/bin/env python3
"""Render each activity's Scenario + Discussion Questions + Debrief to a branded DOCX
(then LibreOffice converts to PDF), plus a combined trainer debrief pack.
"""
import os, sys, subprocess
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
ACT = DOMAIN1 + DOMAIN2 + DOMAIN3
import prodoc


def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "activities")):
            return d
    return os.path.dirname(os.path.dirname(HERE))


REPO = _find_repo(HERE)
ROOT = os.path.join(REPO, "activities")
ASSETS = os.path.join(os.path.dirname(HERE), "assets")
BRAND = RGBColor(0x1F, 0x6F, 0xEB); GREY = RGBColor(0x55, 0x5B, 0x66)
AMBER = RGBColor(0xB4, 0x6B, 0x00); TEAL = RGBColor(0x0E, 0x7A, 0x5E)


def folder_for(a):
    import re
    t = a["title"].split("— ", 1)[-1].lower()
    t = re.sub(r"[^a-z0-9]+", "-", t).strip("-")
    return os.path.join(ROOT, "activity-%02d-%s" % (a["num"], re.sub(r"-+", "-", t)[:48]))


def new_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    prodoc.style_headings(doc)
    return doc


def hdr_block(doc, a):
    p = doc.add_paragraph()
    r = p.add_run("%s  ·  %s" % (C.TITLE, C.COURSE_CODE))
    r.font.size = Pt(9); r.font.color.rgb = GREY
    doc.add_heading("Activity %d — %s" % (a["num"], a["title"].split("— ", 1)[-1]), level=1)
    p = doc.add_paragraph()
    r = p.add_run(a["case_title"]); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BRAND


def meta_table(doc, a):
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for k, v in [("Topic", "Topic %d" % a["topic"]), ("Objective", a["objective"]),
                 ("Duration", a["duration"]), ("Grouping", a["grouping"]),
                 ("Tools", a["services"]),
                 ("Ed-tool", "%s — %s" % (a["edtool"]["name"], a["edtool"]["url"]))]:
        c = t.add_row().cells
        c[0].text = ""; r = c[0].paragraphs[0].add_run(k); r.bold = True; r.font.size = Pt(9.5)
        prodoc._shade_cell(c[0], "E8F0FE")
        c[1].text = ""; c[1].paragraphs[0].add_run(v).font.size = Pt(9.5)
        c[0].width = Inches(1.2); c[1].width = Inches(5.3)
    doc.add_paragraph()


def evidence(doc, a):
    doc.add_heading("The Evidence", level=2)
    p = doc.add_paragraph(); r = p.add_run(a["data"]["caption"]); r.italic = True; r.font.size = Pt(9.5)
    rows = a["data"]["rows"]
    t = doc.add_table(rows=0, cols=len(rows[0])); t.style = "Table Grid"
    cells = t.add_row().cells
    for i, h in enumerate(rows[0]):
        cells[i].text = ""
        r = cells[i].paragraphs[0].add_run(str(h))
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        prodoc._shade_cell(cells[i], "1F6FEB")
    for row in rows[1:]:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(c)); r.font.size = Pt(9)
            if i == 0: r.bold = True
    doc.add_paragraph()


def paras(doc, text, size=10.5):
    for x in [q.strip() for q in text.split("\n\n") if q.strip()]:
        p = doc.add_paragraph(); p.add_run(x).font.size = Pt(size)


def build_activity_pdf(a):
    doc = new_doc()
    hdr_block(doc, a)
    meta_table(doc, a)

    doc.add_heading("The Scenario", level=2)
    paras(doc, a["scenario"], 11)

    evidence(doc, a)

    doc.add_heading("Discussion Questions", level=2)
    for q in a["questions"]:
        doc.add_paragraph(q, style="List Number")

    doc.add_heading("Step-by-Step", level=2)
    for s in a["steps"]:
        doc.add_paragraph(s[0], style="List Number")

    doc.add_heading("The GenAI Prompt", level=2)
    for line in a["prompt"].split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x0B, 0x30, 0x60)
        p.paragraph_format.space_after = Pt(0)

    doc.add_page_break()
    h = doc.add_heading("Trainer Debrief", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Expected answers and the points to draw out. Trainer reference.")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = AMBER
    paras(doc, a["debrief"], 10.5)

    doc.add_heading("Self-Check — Is Your Output Finished?", level=2)
    p = doc.add_paragraph(); p.add_run(a["test"]).font.size = Pt(10.5)

    prodoc.add_page_numbers(doc)
    out = os.path.join(folder_for(a), "debrief.docx")
    doc.save(out)
    return out


made = []
for a in ACT:
    made.append(build_activity_pdf(a))

# ---- combined trainer pack ----
doc = new_doc()
prodoc.add_cover_page(doc, "ACTIVITY & DEBRIEF PACK", C.TITLE, C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_toc(doc)
for a in ACT:
    hdr_block(doc, a)
    meta_table(doc, a)
    doc.add_heading("The Scenario", level=2); paras(doc, a["scenario"], 11)
    evidence(doc, a)
    doc.add_heading("Discussion Questions", level=2)
    for q in a["questions"]: doc.add_paragraph(q, style="List Number")
    doc.add_heading("Trainer Debrief", level=2); paras(doc, a["debrief"], 10.5)
    doc.add_heading("Self-Check", level=2)
    p = doc.add_paragraph(); p.add_run(a["test"]).font.size = Pt(10.5)
    doc.add_page_break()
prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
pack = os.path.join(ROOT, "Activity-Debrief-Pack-%s.docx" % C.SHORT_TITLE)
doc.save(pack)
made.append(pack)
print("built %d docx" % len(made))

# ---- convert all to PDF ----
for f in made:
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", os.path.dirname(f), f],
                   capture_output=True, timeout=300)
    if os.path.exists(f.replace(".docx", ".pdf")):
        os.remove(f)          # keep the folder clean — PDF is the deliverable
print("converted to PDF")
