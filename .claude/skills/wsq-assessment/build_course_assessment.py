#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Generative AI for Problem Solving' (TGS-2023036653):
  - Written Assessment (SAQ) — 8 open-ended KNOWLEDGE questions (K1-K8), aligned to the slides
  - Case Study (CS) — 3 ABILITY tasks (A1/A6, A2/A3, A4/A5/A7), aligned to the activities
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Page 1 is the cover;
page 2 carries Trainee Information + Instructions + Grading; the questions/tasks begin on page 3.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "Generative AI for Problem Solving"
COURSE_CODE = "TGS-2023036653"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = _logo("course-badge.png")   # None if absent → Tertiary-only cover (as LP/LG)

Q_VER, A_VER = "v21", "v21"   # single standardised version across all four files
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.


# ---------------------------------------------------------------- content (single source)
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from assessment_data import WRITTEN, PRACTICAL, SCENARIO

def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → monospace
    code/YAML/command block (indentation preserved); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in code.split("\n"):
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
            rr = b.add_run(ln if ln else " ")
            rr.font.name = "Consolas"; rr.font.size = Pt(9)
            rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
            wt = rr._element.find(qn('w:t'))
            if wt is not None: wt.set(qn('xml:space'), 'preserve')
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))


# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in "
              "the course slides and Learner Guide.", size=10.5, italic=True, color=GREY, after=8)
    per_page = 1 if answers else 2
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
        if i % per_page == 0 and i < len(WRITTEN):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))


def build_cs(answers):
    """Case Study — the instrument on the original TMS paper. Never converted to a PP."""
    doc = base_doc()
    kind = "Case Study (CS) — Answer Key" if answers else "Case Study (CS)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Case Study Assessment" if answers else "Case Study Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has completed all case study tasks and demonstrated the ability to identify "
                     "root causes, evaluate solutions and plan and evaluate their implementation.")
        page_break(doc)
    para(doc, "Case Study", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Read the scenario below and complete all three tasks.", size=10.5, italic=True,
         color=GREY, after=6)
    para(doc, "Scenario: Lumina Home", size=11.5, bold=True, after=3)
    for blk in SCENARIO.split("\n\n"):
        para(doc, blk, size=11, after=6)
    page_break(doc)
    for i, (label, crit, prompt, cap, pts) in enumerate(PRACTICAL, 1):
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=3, before=4)
        for blk in prompt.split("\n\n"):
            para(doc, blk, size=11, after=4)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, lines=pts if answers else None, height_pt=170)
        if i < len(PRACTICAL):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to Case Study - {TITLE} - {suffix}.docx" if answers
            else f"Case Study - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))


def check_coverage():
    """Fail the build on an orphan K or A — a missing code is never shipped."""
    ks = [c for c, *_ in WRITTEN]
    want_k = ["K%d" % i for i in range(1, 9)]
    missing_k = [k for k in want_k if k not in ks]
    acodes = set()
    for _, crit, *_ in PRACTICAL:
        for a in crit.replace("/", " ").split():
            acodes.add(a.strip())
    want_a = ["A%d" % i for i in range(1, 8)]
    missing_a = [a for a in want_a if a not in acodes]
    print("\n  K coverage (WA):", ", ".join("%s->Q%d" % (c, i) for i, (c, *_) in enumerate(WRITTEN, 1)))
    print("  A coverage (CS):", ", ".join("%s->%s" % (lbl, crit) for lbl, crit, *_ in PRACTICAL))
    if missing_k or missing_a:
        raise SystemExit("COVERAGE FAILURE — missing K: %s  missing A: %s" % (missing_k, missing_a))
    print("  coverage OK — all K1-K8 and A1-A7 covered.")


if __name__ == "__main__":
    os.makedirs(OUT, exist_ok=True)
    print("Building WSQ assessment set...")
    build_wa(answers=False); build_wa(answers=True)
    build_cs(answers=False); build_cs(answers=True)
    check_coverage()
    print(f"Done. WA: {len(WRITTEN)} questions - CS: {len(PRACTICAL)} tasks.")
